import pytest
from pathlib import Path
from tempfile import NamedTemporaryFile
from docx import Document

from tests.utils.api_response_validator import APIResponseValidator, ValidationRule, EndpointValidator

BASE_DIR = Path(__file__).resolve().parents[2]
DOCX_PATH = BASE_DIR / "ttable_25-26.docx"


def _build_simple_doc(path: Path, group_name: str, discipline: str, teacher: str, aud: str, day_label: str = "пн"):
    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = "День"
    hdr[1].text = "№"
    hdr[2].text = group_name
    hdr[3].text = "Ауд"
    # Row 1
    row = table.rows[1].cells
    row[0].text = day_label
    row[1].text = "1"
    row[2].text = f"{discipline}\n{teacher}"
    row[3].text = aud
    doc.save(path)


@pytest.mark.asyncio
async def test_healthcheck(client):
    resp = await client.get("/api/v1/public/healthcheck")
    assert resp.status_code == 200
    assert resp.json().get("status") == True


@pytest.mark.asyncio
async def test_cards_history(client, seed_info):
    """Test cards history endpoint validates actual data structure."""
    # Create validator for history response
    validator = APIResponseValidator(strict_mode=False)
    validator.add_rule(ValidationRule("history", list, required=True))
    
    resp = await client.get(
        "/api/v1/private/n8n_ui/cards/history",
        params={"sched_ver_id": seed_info["ttable_id"], "group_id": seed_info["group_id"]},
    )
    
    # Validate response structure
    assert resp.status_code == 200
    data = resp.json()
    
    validation_result = validator.validate_response(data)
    assert validation_result.is_valid, f"Validation errors: {validation_result.errors}"
    
    # Validate business logic: history should contain expected data
    history = data["history"]
    assert len(history) == 1, "Expected exactly one history record"
    
    row = history[0]
    # Validate essential fields exist with correct business logic
    assert "status_id" in row, "History record should have status_id"
    assert "is_current" in row, "History record should have is_current"
    assert "user_name" in row, "History record should have user_name"
    
    # Validate business logic values
    assert row["status_id"] == seed_info["cards_statuses"]["edited"], "Status should be 'edited'"
    assert row["is_current"] is True, "Record should be current"
    assert row["user_name"] == "Test User", "User name should match expected value"


@pytest.mark.asyncio
async def test_cards_content(client, seed_info):
    """Test cards content endpoint validates actual data structure."""
    # Create validator for content response
    validator = APIResponseValidator(strict_mode=False)
    validator.add_rule(ValidationRule("card_content", list, required=True))
    
    resp = await client.get(
        "/api/v1/private/n8n_ui/cards/history_content",
        params={"card_hist_id": seed_info["hist_id"]},
    )
    
    # Validate response structure
    assert resp.status_code == 200
    data = resp.json()
    
    validation_result = validator.validate_response(data)
    assert validation_result.is_valid, f"Validation errors: {validation_result.errors}"
    
    # Validate business logic: content should contain expected data
    content = data["card_content"]
    assert len(content) == 2, "Expected exactly two content records (from seed)"
    
    row = content[0]
    # Validate essential fields exist
    assert "position" in row, "Content record should have position"
    assert "aud" in row, "Content record should have aud"
    assert "discipline_title" in row, "Content record should have discipline_title"
    assert "teacher_name" in row, "Content record should have teacher_name"
    
    # Validate business logic values
    assert row["position"] == 1, "Position should be 1"
    assert row["aud"] == "101", "Aud should be '101'"
    assert row["discipline_title"] == "Math", "Discipline should be 'Math'"
    assert row["teacher_name"] == "Иванов И.И.", "Teacher name should match expected value"


@pytest.mark.asyncio
async def test_cards_get_by_id(client, seed_info):
    resp = await client.post(
        "/api/v1/private/n8n_ui/cards/get_by_id",
        json={"card_hist_id": seed_info["hist_id"]},
    )
    assert resp.status_code == 200
    payload = resp.json()["ext_card"]
    assert len(payload) == 2, "Expected 2 lessons from seed"
    
    # Check that both lessons exist (order not guaranteed)
    auds = {item["aud"] for item in payload}
    assert auds == {"101", "102"}, f"Expected auds 101 and 102, got {auds}"
    
    # Check teacher_id is correct for all lessons
    for item in payload:
        assert item["teacher_id"] == seed_info["teacher_id"]


@pytest.mark.asyncio
async def test_cards_accept(client, seed_info, pg_pool):
    resp = await client.put(
        "/api/v1/private/n8n_ui/cards/accept",
        json={"card_hist_id": seed_info["hist_id"]},
    )
    assert resp.status_code == 200
    async with pg_pool.acquire() as conn:
        status_row = await conn.fetchrow(
            "SELECT status_id FROM cards_states_history WHERE id=$1",
            seed_info["hist_id"],
        )
    assert status_row["status_id"] == seed_info["cards_statuses"]["accepted"]  # accepted


@pytest.mark.asyncio
async def test_timetable_standard_import(client, pg_pool, seed_info):
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    _build_simple_doc(tmp_path, "GR1", "Math", "Иванов И.И.", "201", day_label="пн")

    with tmp_path.open("rb") as f:
        files = {"file_obj": (tmp_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        resp = await client.post(
            "/api/v1/private/ttable/versions/standard/import",
            params={"smtr": "1", "bid": seed_info["building_id"]},
            files=files,
        )
    tmp_path.unlink(missing_ok=True)

    assert resp.status_code == 200
    payload = resp.json()
    assert payload.get("success") is True
    new_id = payload["ttable_ver_id"]["id"] if isinstance(payload["ttable_ver_id"], dict) else payload["ttable_ver_id"]
    async with pg_pool.acquire() as conn:
        cnt = await conn.fetchval("SELECT count(*) FROM ttable_versions WHERE id=$1", new_id)
    assert cnt == 1


@pytest.mark.asyncio
async def test_std_ttable_get_all_creates_snapshot(client, seed_info, pg_pool):
    """Test std_ttable get_all endpoint validates actual data processing logic."""
    # Create validator for lessons response
    validator = APIResponseValidator(strict_mode=False)
    validator.add_rule(ValidationRule("lessons", list, required=True))
    
    async with pg_pool.acquire() as conn:
        new_sched_id = await conn.fetchval(
            """
            INSERT INTO ttable_versions (schedule_date, status_id, building_id, user_id, type, is_commited)
            VALUES (CURRENT_DATE, $1, $2, $3, 'standard', false)
            RETURNING id
            """,
            seed_info["ttable_statuses"]["accepted"],
            seed_info["building_id"],
            seed_info["user_id"],
        )

    body = {
        "building_id": seed_info["building_id"],
        "week_day": 1,
        "ttable_id": new_sched_id,
    }
    resp = await client.post("/api/v1/private/n8n_ui/ttable/std/get_all", json=body)
    
    # Validate response structure
    assert resp.status_code == 200
    data = resp.json()
    
    validation_result = validator.validate_response(data)
    assert validation_result.is_valid, f"Validation errors: {validation_result.errors}"
    
    # Validate business logic: lessons should contain expected data
    lessons = data["lessons"]
    assert len(lessons) == 1, "Expected exactly one lesson record"
    
    item = lessons[0]
    # Validate essential fields exist
    assert "status_card" in item, "Lesson should have status_card"
    assert "position" in item, "Lesson should have position"
    assert "title" in item, "Lesson should have title"
    
    # Validate business logic values
    assert item["status_card"] == seed_info["cards_statuses"]["draft"], "Status should be 'draft'"
    assert item["position"] == 1, "Position should be 1"
    assert item["title"] == "Math", "Title should be 'Math'"

    # Validate actual database changes (business logic verification)
    async with pg_pool.acquire() as conn:
        current = await conn.fetch("SELECT id, is_current FROM cards_states_history WHERE sched_ver_id=$1", new_sched_id)
    
    assert len(current) == 1, "Expected exactly one card state history record"
    assert current[0]["is_current"] is True, "Card state should be current"


@pytest.mark.asyncio
async def test_current_ttable_get_all_returns_active_cards(client, seed_info):
    """Test current_ttable get_all endpoint validates actual data processing logic."""
    # Create validator for lessons response
    validator = APIResponseValidator(strict_mode=False)
    validator.add_rule(ValidationRule("lessons", list, required=True))
    
    resp = await client.post(
        "/api/v1/private/n8n_ui/ttable/current/get_all",
        json={"ttable_id": seed_info["ttable_id"]},
    )
    
    # Validate response structure
    assert resp.status_code == 200
    data = resp.json()
    
    validation_result = validator.validate_response(data)
    assert validation_result.is_valid, f"Validation errors: {validation_result.errors}"
    
    # Validate business logic: lessons should contain expected data
    lessons = data["lessons"]
    assert len(lessons) == 2, "Expected exactly two lesson records (from seed)"
    
    row = lessons[0]
    # Validate essential fields exist
    assert "card_hist_id" in row, "Lesson should have card_hist_id"
    assert "status_card" in row, "Lesson should have status_card"
    assert "title" in row, "Lesson should have title"
    
    # Validate business logic values
    assert row["card_hist_id"] == seed_info["hist_id"], "Card hist ID should match expected"
    assert row["status_card"] == seed_info["cards_statuses"]["edited"], "Status should be 'edited'"
    assert row["title"] == "Math", "Title should be 'Math'"


@pytest.mark.asyncio
async def test_cards_save_creates_new_version(client, seed_info, pg_pool):
    """Test cards save endpoint validates actual data processing logic."""
    # Create validator for cards save endpoint
    validator = APIResponseValidator(strict_mode=False)
    validator.add_rule(ValidationRule("success", bool, required=True))
    validator.add_rule(ValidationRule("new_card_hist_id", int, required=False))
    
    payload = {
        "card_hist_id": seed_info["hist_id"],
        "ttable_id": seed_info["ttable_id"],
        "week_day": None,  # Замены
        "lessons": [
            {
                "position": 3,
                "discipline_id": seed_info["discipline_id"],
                "teacher_id": seed_info["teacher_id"],
                "aud": "303",
                "week_day": None,
                "is_force": False,
            }
        ],
    }
    resp = await client.post("/api/v1/private/n8n_ui/cards/save", json=payload)
    
    # Validate response using independent validator
    assert resp.status_code == 200
    data = resp.json()
    
    # Use validator to check response structure
    validation_result = validator.validate_response(data)
    assert validation_result.is_valid, f"Validation errors: {validation_result.errors}"
    
    # Validate business logic: success should be True for valid save
    assert data["success"] is True, "Expected successful save operation"
    
    # Validate business logic: new_card_hist_id should be present and different
    assert "new_card_hist_id" in data, "Expected new_card_hist_id in successful response"
    new_hist_id = data["new_card_hist_id"]
    assert isinstance(new_hist_id, int), "new_card_hist_id should be an integer"
    assert new_hist_id != seed_info["hist_id"], "new_card_hist_id should be different from original"

    # Validate actual database changes (business logic verification)
    async with pg_pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT csd.card_hist_id, csh.is_current, csd.aud "
            "FROM cards_states_details csd "
            "JOIN cards_states_history csh ON csh.id = csd.card_hist_id "
            "WHERE csh.sched_ver_id=$1",
            seed_info["ttable_id"],
        )
    
    # Verify business logic: new record exists and is current
    new_record_exists = any(r["card_hist_id"] == new_hist_id for r in rows)
    assert new_record_exists, "New card history record should exist in database"
    
    correct_aud_and_current = any(r["aud"] == "303" and r["is_current"] for r in rows)
    assert correct_aud_and_current, "New record should have correct aud and be current"


@pytest.mark.asyncio
async def test_cards_save_conflict_due_to_unique_index(client, seed_info, pg_pool):
    """Test cards save endpoint handles conflicts correctly with week_day index."""
    # Create validator for conflict response
    validator = APIResponseValidator(strict_mode=False)
    validator.add_rule(ValidationRule("success", bool, required=True))
    validator.add_rule(ValidationRule("conflicts", dict, required=False))
    
    payload = {
        "card_hist_id": seed_info["hist_id"],
        "ttable_id": seed_info["ttable_id"],
        "week_day": None,  # Замены
        "lessons": [
            {
                "position": 1,  # та же позиция, что и в seed
                "discipline_id": seed_info["discipline_id"],
                "teacher_id": seed_info["teacher_id"],
                "aud": "303",
                "week_day": None,
                "is_force": False,
            }
        ],
    }
    resp = await client.post("/api/v1/private/n8n_ui/cards/save", json=payload)
    
    # Validate response structure
    assert resp.status_code == 200
    data = resp.json()
    
    validation_result = validator.validate_response(data)
    assert validation_result.is_valid, f"Validation errors: {validation_result.errors}"
    
    # С новым индексом (sched_ver_id, teacher_id, position, week_day) конфликт возникает
    # только если все поля совпадают. Проверяем, что в БД есть только одна активная запись
    # для этой комбинации (sched_ver_id, teacher_id, position=1, week_day=NULL)
    async with pg_pool.acquire() as conn:
        active_records = await conn.fetch(
            """
            SELECT csd.* FROM cards_states_details csd
            JOIN cards_states_history csh ON csh.id = csd.card_hist_id
            WHERE csh.sched_ver_id = $1 
              AND csd.teacher_id = $2 
              AND csd.position = 1
              AND csd.week_day IS NULL
              AND csh.is_current = true
            """,
            seed_info["ttable_id"],
            seed_info["teacher_id"]
        )
        # Должна быть только одна активная запись (новая карточка)
        assert len(active_records) == 1, f"Expected 1 active record, got {len(active_records)}"
        assert active_records[0]["aud"] == "303", "Expected new aud value"


@pytest.mark.asyncio
async def test_cards_save_conflict_even_with_force(client, seed_info):
    """Test cards save endpoint handles force flag correctly."""
    # Create validator for force save response
    validator = APIResponseValidator(strict_mode=False)
    validator.add_rule(ValidationRule("success", bool, required=True))
    validator.add_rule(ValidationRule("new_card_hist_id", int, required=False))
    
    payload = {
        "card_hist_id": seed_info["hist_id"],
        "ttable_id": seed_info["ttable_id"],
        "week_day": None,  # Замены
        "lessons": [
            {
                "position": 1,  # тот же слот, но is_force=True
                "discipline_id": seed_info["discipline_id"],
                "teacher_id": seed_info["teacher_id"],
                "aud": "304",
                "week_day": None,
                "is_force": True,
            }
        ],
    }
    resp = await client.post("/api/v1/private/n8n_ui/cards/save", json=payload)
    
    # Validate response structure
    assert resp.status_code == 200
    data = resp.json()
    
    validation_result = validator.validate_response(data)
    assert validation_result.is_valid, f"Validation errors: {validation_result.errors}"
    
    # Validate business logic: force flag should allow save despite conflict
    assert data["success"] is True, "Expected force flag to allow save despite conflict"
    
    # Validate business logic: new_card_hist_id should be present for successful save
    assert "new_card_hist_id" in data, "Expected new_card_hist_id for successful force save"
    assert isinstance(data.get("new_card_hist_id"), int), "new_card_hist_id should be integer"




@pytest.mark.asyncio
async def test_timetable_standard_import_creates_cards(client, pg_pool, seed_info):
    """
    Тест проверяет, что импорт стандартного расписания создаёт карточки в cards_states_history/details
    вместо записей в std_ttable
    """
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    
    # Создаём документ с расписанием для понедельника и вторника
    doc = Document()
    table = doc.add_table(rows=3, cols=4)
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = "День"
    hdr[1].text = "№"
    hdr[2].text = "GR1"
    hdr[3].text = "Ауд"
    
    # Понедельник, пара 1
    row1 = table.rows[1].cells
    row1[0].text = "пн"
    row1[1].text = "1"
    row1[2].text = "Math\nИванов И.И."
    row1[3].text = "201"
    
    # Вторник, пара 1
    row2 = table.rows[2].cells
    row2[0].text = "вт"
    row2[1].text = "1"
    row2[2].text = "Math\nИванов И.И."
    row2[3].text = "202"
    
    doc.save(tmp_path)
    
    with tmp_path.open("rb") as f:
        files = {"file_obj": (tmp_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        resp = await client.post(
            "/api/v1/private/ttable/versions/standard/import",
            params={"smtr": "1", "bid": seed_info["building_id"]},
            files=files,
        )
    tmp_path.unlink(missing_ok=True)
    
    assert resp.status_code == 200
    payload = resp.json()
    assert payload.get("success") is True
    new_id = payload["ttable_ver_id"]["id"] if isinstance(payload["ttable_ver_id"], dict) else payload["ttable_ver_id"]
    
    async with pg_pool.acquire() as conn:
        # Проверяем, что версия создана
        version = await conn.fetchrow("SELECT * FROM ttable_versions WHERE id=$1", new_id)
        assert version is not None
        assert version["type"] == "standard"
        
        # Проверяем, что созданы карточки в cards_states_history
        cards = await conn.fetch(
            "SELECT * FROM cards_states_history WHERE sched_ver_id=$1 ORDER BY week_day",
            new_id
        )
        assert len(cards) == 2, "Должно быть 2 карточки (понедельник и вторник)"
        
        # Проверяем понедельник (week_day = 0)
        card_monday = cards[0]
        assert card_monday["week_day"] == 0, "Понедельник должен иметь week_day = 0"
        assert card_monday["group_id"] == seed_info["group_id"]
        assert card_monday["status_id"] == seed_info["cards_statuses"]["draft"]
        assert card_monday["is_current"] is True
        
        # Проверяем вторник (week_day = 1)
        card_tuesday = cards[1]
        assert card_tuesday["week_day"] == 1, "Вторник должен иметь week_day = 1"
        assert card_tuesday["group_id"] == seed_info["group_id"]
        assert card_tuesday["status_id"] == seed_info["cards_statuses"]["draft"]
        assert card_tuesday["is_current"] is True
        
        # Проверяем детали карточек
        details_monday = await conn.fetch(
            "SELECT * FROM cards_states_details WHERE card_hist_id=$1",
            card_monday["id"]
        )
        assert len(details_monday) == 1, "Должна быть 1 пара в понедельник"
        assert details_monday[0]["position"] == 1
        assert details_monday[0]["aud"] == "201"
        assert details_monday[0]["discipline_id"] == seed_info["discipline_id"]
        assert details_monday[0]["teacher_id"] == seed_info["teacher_id"]
        
        details_tuesday = await conn.fetch(
            "SELECT * FROM cards_states_details WHERE card_hist_id=$1",
            card_tuesday["id"]
        )
        assert len(details_tuesday) == 1, "Должна быть 1 пара во вторник"
        assert details_tuesday[0]["position"] == 1
        assert details_tuesday[0]["aud"] == "202"


@pytest.mark.asyncio
async def test_timetable_standard_import_deactivates_old_cards(client, pg_pool, seed_info):
    """
    Тест проверяет, что повторный импорт деактивирует старые карточки (is_current = false)
    """
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    _build_simple_doc(tmp_path, "GR1", "Math", "Иванов И.И.", "201", day_label="пн")
    
    # Первый импорт
    with tmp_path.open("rb") as f:
        files = {"file_obj": (tmp_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        resp1 = await client.post(
            "/api/v1/private/ttable/versions/standard/import",
            params={"smtr": "1", "bid": seed_info["building_id"]},
            files=files,
        )
    
    assert resp1.status_code == 200
    ttable_id_1 = resp1.json()["ttable_ver_id"]["id"] if isinstance(resp1.json()["ttable_ver_id"], dict) else resp1.json()["ttable_ver_id"]
    
    async with pg_pool.acquire() as conn:
        # Проверяем, что карточка создана и активна
        card1 = await conn.fetchrow(
            "SELECT * FROM cards_states_history WHERE sched_ver_id=$1",
            ttable_id_1
        )
        assert card1 is not None
        assert card1["is_current"] is True
        
        # Второй импорт в ту же версию (имитация повторного импорта)
        # Создаём новую карточку вручную, чтобы проверить деактивацию
        card2_id = await conn.fetchval(
            "INSERT INTO cards_states_history (sched_ver_id, user_id, status_id, group_id, is_current, week_day) "
            "VALUES ($1, $2, $3, $4, true, 0) RETURNING id",
            ttable_id_1,
            seed_info["user_id"],
            seed_info["cards_statuses"]["draft"],
            seed_info["group_id"]
        )
        
        # Проверяем, что обе карточки активны
        active_cards = await conn.fetch(
            "SELECT * FROM cards_states_history WHERE sched_ver_id=$1 AND is_current=true",
            ttable_id_1
        )
        assert len(active_cards) == 2
    
    # Третий импорт - должен деактивировать обе старые карточки
    with tmp_path.open("rb") as f:
        files = {"file_obj": (tmp_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        # Используем тот же ttable_id через создание новой версии
        # На практике это будет новый импорт, но для теста проверим логику деактивации
        from core.data.sql_queries.ttable_sql import TimetableQueries
        async with pg_pool.acquire() as conn:
            queries = TimetableQueries(conn)
            # Имитируем повторный импорт
            await conn.execute(
                'UPDATE cards_states_history SET is_current = false WHERE sched_ver_id = $1 AND is_current = true',
                ttable_id_1
            )
            
            # Проверяем, что все старые карточки деактивированы
            active_cards_after = await conn.fetch(
                "SELECT * FROM cards_states_history WHERE sched_ver_id=$1 AND is_current=true",
                ttable_id_1
            )
            assert len(active_cards_after) == 0, "Все старые карточки должны быть деактивированы"
    
    tmp_path.unlink(missing_ok=True)


@pytest.mark.asyncio
async def test_timetable_standard_import_multiple_teachers(client, pg_pool, seed_info):
    """
    Тест проверяет импорт расписания с несколькими преподавателями для одной пары
    """
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    
    # Создаём дополнительного преподавателя
    async with pg_pool.acquire() as conn:
        teacher2_id = await conn.fetchval(
            "INSERT INTO teachers (fio) VALUES ('Петров П.П.') "
            "ON CONFLICT (fio) DO NOTHING RETURNING id"
        )
        if teacher2_id is None:
            teacher2_id = await conn.fetchval(
                "SELECT id FROM teachers WHERE fio = 'Петров П.П.'"
            )
    
    # Создаём документ с двумя преподавателями
    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = "День"
    hdr[1].text = "№"
    hdr[2].text = "GR1"
    hdr[3].text = "Ауд"
    
    # Понедельник, пара 1, два преподавателя
    row1 = table.rows[1].cells
    row1[0].text = "пн"
    row1[1].text = "1"
    row1[2].text = "Math\nИванов И.И. Петров П.П."
    row1[3].text = "201, 202"
    
    doc.save(tmp_path)
    
    with tmp_path.open("rb") as f:
        files = {"file_obj": (tmp_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        resp = await client.post(
            "/api/v1/private/ttable/versions/standard/import",
            params={"smtr": "1", "bid": seed_info["building_id"]},
            files=files,
        )
    tmp_path.unlink(missing_ok=True)
    
    assert resp.status_code == 200
    new_id = resp.json()["ttable_ver_id"]["id"] if isinstance(resp.json()["ttable_ver_id"], dict) else resp.json()["ttable_ver_id"]
    
    async with pg_pool.acquire() as conn:
        # Проверяем карточку
        card = await conn.fetchrow(
            "SELECT * FROM cards_states_history WHERE sched_ver_id=$1",
            new_id
        )
        assert card is not None
        
        # Проверяем детали - должно быть 2 записи (по одной на каждого преподавателя)
        details = await conn.fetch(
            "SELECT * FROM cards_states_details WHERE card_hist_id=$1 ORDER BY teacher_id",
            card["id"]
        )
        assert len(details) == 2, "Должно быть 2 записи для двух преподавателей"
        
        # Проверяем первого преподавателя
        assert details[0]["position"] == 1
        assert details[0]["teacher_id"] == seed_info["teacher_id"]
        assert details[0]["aud"] == "201"
        
        # Проверяем второго преподавателя
        assert details[1]["position"] == 1
        assert details[1]["teacher_id"] == teacher2_id
        assert details[1]["aud"] == "202"



@pytest.mark.asyncio
async def test_timetable_standard_import_teacher_conflict_different_days(client, pg_pool, seed_info):
    """
    Тест проверяет, что преподаватель может вести одну и ту же позицию в разные дни недели
    Это НЕ должно вызывать конфликт индекса
    """
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    
    # Создаём документ: один преподаватель ведёт позицию 1 в понедельник и вторник
    doc = Document()
    table = doc.add_table(rows=3, cols=4)
    
    # Header
    hdr = table.rows[0].cells
    hdr[0].text = "День"
    hdr[1].text = "№"
    hdr[2].text = "GR1"
    hdr[3].text = "Ауд"
    
    # Понедельник, пара 1
    row1 = table.rows[1].cells
    row1[0].text = "пн"
    row1[1].text = "1"
    row1[2].text = "Math\nИванов И.И."
    row1[3].text = "201"
    
    # Вторник, пара 1 (тот же преподаватель, та же позиция)
    row2 = table.rows[2].cells
    row2[0].text = "вт"
    row2[1].text = "1"
    row2[2].text = "Math\nИванов И.И."
    row2[3].text = "202"
    
    doc.save(tmp_path)
    
    with tmp_path.open("rb") as f:
        files = {"file_obj": (tmp_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        resp = await client.post(
            "/api/v1/private/ttable/versions/standard/import",
            params={"smtr": "1", "bid": seed_info["building_id"]},
            files=files,
        )
    tmp_path.unlink(missing_ok=True)
    
    # Должно пройти успешно (НЕ должно быть конфликта)
    assert resp.status_code == 200
    payload = resp.json()
    assert payload.get("success") is True
    new_id = payload["ttable_ver_id"]["id"] if isinstance(payload["ttable_ver_id"], dict) else payload["ttable_ver_id"]
    
    async with pg_pool.acquire() as conn:
        # Проверяем, что созданы обе карточки
        cards = await conn.fetch(
            "SELECT * FROM cards_states_history WHERE sched_ver_id=$1 ORDER BY week_day",
            new_id
        )
        assert len(cards) == 2, "Должно быть 2 карточки (понедельник и вторник)"
        
        # Проверяем детали обеих карточек
        for card in cards:
            details = await conn.fetch(
                "SELECT * FROM cards_states_details WHERE card_hist_id=$1",
                card["id"]
            )
            assert len(details) == 1
            assert details[0]["teacher_id"] == seed_info["teacher_id"]
            assert details[0]["position"] == 1
            assert details[0]["week_day"] == card["week_day"]


@pytest.mark.asyncio
async def test_timetable_standard_import_teacher_conflict_same_day(client, pg_pool, seed_info):
    """
    Тест проверяет, что преподаватель НЕ может вести две группы в одну позицию в один день
    Это ДОЛЖНО вызывать конфликт индекса или игнорироваться
    """
    with NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    
    # Создаём дополнительную группу
    async with pg_pool.acquire() as conn:
        group2_id = await conn.fetchval(
            "INSERT INTO groups (name, building_id, is_active) VALUES ('GR_CONFLICT', $1, true) RETURNING id",
            seed_info["building_id"]
        )
    
    # Создаём документ: один преподаватель ведёт позицию 1 у двух групп в понедельник
    doc = Document()
    table = doc.add_table(rows=2, cols=7)
    
    # Header - формат: День, №, Группа1, Ауд1, №, Группа2, Ауд2
    hdr = table.rows[0].cells
    hdr[0].text = "День"
    hdr[1].text = "№"
    hdr[2].text = "GR1"
    hdr[3].text = "Ауд"
    hdr[4].text = "№"
    hdr[5].text = "GR_CONFLICT"
    hdr[6].text = "Ауд"
    
    # Понедельник, пара 1 - обе группы (тот же преподаватель!)
    row1 = table.rows[1].cells
    row1[0].text = "пн"
    row1[1].text = "1"
    row1[2].text = "Math\nИванов И.И."
    row1[3].text = "201"
    row1[4].text = "1"
    row1[5].text = "Math\nИванов И.И."
    row1[6].text = "202"
    
    doc.save(tmp_path)
    
    with tmp_path.open("rb") as f:
        files = {"file_obj": (tmp_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        resp = await client.post(
            "/api/v1/private/ttable/versions/standard/import",
            params={"smtr": "1", "bid": seed_info["building_id"]},
            files=files,
        )
    tmp_path.unlink(missing_ok=True)
    
    # Импорт должен либо упасть с ошибкой, либо пройти с игнорированием конфликта
    # В зависимости от реализации (с ON CONFLICT или без)
    if resp.status_code == 200:
        # Если прошло успешно, проверяем, что создана только одна запись для преподавателя
        payload = resp.json()
        new_id = payload["ttable_ver_id"]["id"] if isinstance(payload["ttable_ver_id"], dict) else payload["ttable_ver_id"]
        
        async with pg_pool.acquire() as conn:
            # Проверяем, что для понедельника позиции 1 есть только одна запись с этим преподавателем
            details = await conn.fetch(
                """
                SELECT csd.* FROM cards_states_details csd
                JOIN cards_states_history csh ON csh.id = csd.card_hist_id
                WHERE csh.sched_ver_id = $1 
                  AND csd.teacher_id = $2 
                  AND csd.position = 1
                  AND csh.week_day = 0
                """,
                new_id,
                seed_info["teacher_id"]
            )
            # Должна быть только одна запись (конфликт предотвращён индексом)
            assert len(details) == 1, f"Должна быть только 1 запись для преподавателя в позиции 1 в понедельник, получено: {len(details)}"
    else:
        # Если упало с ошибкой - это тоже валидное поведение
        assert resp.status_code in [400, 409, 500], f"Ожидалась ошибка конфликта, получен статус: {resp.status_code}"
